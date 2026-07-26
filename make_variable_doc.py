from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.append("/Users/jhonpatino/.codex/plugins/cache/openai-primary-runtime/documents/26.723.12215/skills/documents/scripts")
from table_geometry import apply_table_geometry  # type: ignore


OUT = Path("/Users/jhonpatino/Downloads/TABLAS SIKA/Explicacion_variables_dashboard.docx")


def set_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, *, bold=False, size=10.5, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)


def format_body_paragraph(paragraph, before=0, after=6, line_spacing=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_body_paragraph(p, before=0, after=2, line_spacing=1.0)
    run = p.add_run("Explicación de variables del dashboard")
    set_font(run, name="Calibri", size=22, color="0B2545", bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_body_paragraph(p, before=0, after=10, line_spacing=1.0)
    run = p.add_run("Cómo se tomaron, por qué existen y cómo se usan en la arquitectura modular")
    set_font(run, name="Calibri", size=11.5, color="555555", italic=True)

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    apply_table_geometry(meta, [1700, 7660], table_width_dxa=9360, indent_dxa=120)
    meta_rows = [
        ("Documento", "Guía de variables y decisiones de diseño"),
        ("Contexto", "Dashboard Streamlit modular con CSV en DEV y Access en PROD"),
        ("Alcance", "Explicación de configuración, limpieza, negocio y UI"),
    ]
    for row_idx, (label, value) in enumerate(meta_rows):
        left = meta.rows[row_idx].cells[0]
        right = meta.rows[row_idx].cells[1]
        shade_cell(left, "F2F4F7")
        set_cell_text(left, label, bold=True, size=10.5, color="0B2545")
        set_cell_text(right, value, size=10.5)

    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style = f"Heading {level}"
    p.style = style
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=14.5, color="0B2545", bold=True)
    else:
        set_font(run, size=12, color="1F4D78", bold=True)
    return p


def add_body(doc, text, *, bold_lead=None):
    p = doc.add_paragraph()
    format_body_paragraph(p)
    if bold_lead and text.startswith(bold_lead):
        lead, rest = text.split(":", 1)
        r1 = p.add_run(f"{lead}:")
        set_font(r1, size=11, bold=True)
        r2 = p.add_run(rest)
        set_font(r2, size=11)
    else:
        run = p.add_run(text)
        set_font(run, size=11)


def build_table(doc, title, headers, rows, widths, header_fill="E8EEF5"):
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        shade_cell(hdr_cells[idx], header_fill)
        set_cell_text(hdr_cells[idx], header, bold=True, size=10.2, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT
            if idx == 0 and len(headers) > 2:
                align = WD_ALIGN_PARAGRAPH.LEFT
            elif isinstance(value, (int, float)) or str(value).replace(".", "", 1).isdigit():
                align = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], str(value), size=10, align=align)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [("Heading 1", 14.5, "0B2545"), ("Heading 2", 12, "1F4D78"), ("Heading 3", 11, "1F4D78")]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    core = doc.core_properties
    core.title = "Explicación de variables del dashboard"
    core.subject = "Documento de variables y criterios de diseño"
    core.author = "OpenAI"

    add_title_block(doc)

    add_heading(doc, "1. Cómo tomé las variables")
    add_body(doc, "No tomé todas las columnas como variables de negocio. Separé primero las que configuran la app, luego las que salen del procesamiento, y por último las que solo existen para la interfaz.")
    add_body(doc, "La regla fue simple: si una variable cambia la fuente de datos o la conexión, pertenece a configuración; si nace de limpiar o combinar datos, pertenece al modelo; si solo sirve para pintar filtros o indicadores, pertenece a UI.")
    add_body(doc, "Con eso evité mezclar consultas, transformaciones y presentación dentro de una sola capa.")

    config_rows = [
        ("mode", "Define si la app trabaja en DEV o PROD", "Permite cambiar de CSV a Access sin tocar la lógica de negocio"),
        ("csv_paths", "Mapa de tablas CSV", "Centraliza las rutas de desarrollo"),
        ("access_db_path", "Ruta del archivo .accdb", "Hace posible la conexión directa en producción"),
        ("access_connection_string", "Cadena ODBC lista para pyodbc", "Evita construir la conexión en varios sitios"),
        ("access_tables", "Nombres reales de tablas Access", "Desacopla el nombre lógico de la app del nombre físico de la base"),
    ]
    build_table(
        doc,
        "2. Variables de configuración",
        ["Variable", "Qué representa", "Por qué la elegí"],
        config_rows,
        [1600, 3150, 4610],
    )

    derived_rows = [
        ("formula_key", "Normalización de Recipe1Name / Formula", "Permite cruces estables entre tablas aunque cambien mayúsculas, espacios o acentos"),
        ("report_datetime", "ReportDate + ReportTime", "Unifica fecha y hora en una sola marca temporal"),
        ("report_day", "Solo la fecha", "Facilita filtros diarios y series temporales"),
        ("total_target_components", "Suma de targets de silos", "Resume el objetivo total de la mezcla"),
        ("total_real_components", "Suma de reales de silos", "Resume el consumo/producción real"),
        ("total_difference", "Suma de diferencias", "Da una señal global del desvío"),
        ("out_of_tolerance_count", "Conteo de componentes fuera de tolerancia", "Convierte muchas columnas binarias en una sola métrica"),
        ("out_of_tolerance_flag", "Indicador 0/1", "Sirve para KPI y porcentaje de alertas"),
        ("snapshot_datetime", "Fecha + hora de BSton", "Ordena la tabla operativa en el tiempo"),
        ("display_name", "Mejor nombre disponible de la fórmula", "Muestra un texto legible al usuario sin perder la llave interna"),
    ]
    build_table(
        doc,
        "3. Variables derivadas de limpieza y negocio",
        ["Variable", "De dónde sale", "Para qué sirve"],
        derived_rows,
        [1850, 3200, 4310],
    )

    ui_rows = [
        ("formula_mapping", "Diccionario display_name -> formula_key", "El usuario ve un nombre limpio y el filtro usa la llave estable"),
        ("filters", "Objeto SidebarFilters", "Agrupa fecha, fórmulas y operadores en una sola salida"),
        ("source_fingerprint", "Timestamps de origen", "Invalida la caché cuando cambian los archivos"),
        ("total_records", "len(filtered.m1)", "KPI de volumen"),
        ("yield_pct", "real / target", "KPI rápido de cumplimiento"),
        ("out_rate", "Promedio de out_of_tolerance_flag", "KPI de desviación"),
    ]
    build_table(
        doc,
        "4. Variables de interfaz y orquestación",
        ["Variable", "Origen", "Uso"],
        ui_rows,
        [2200, 3000, 4160],
    )

    add_heading(doc, "5. Por qué lo hice así")
    reasons = [
        ("Separación de capas", "La UI no debe conocer detalles de SQL ni de lectura de archivos."),
        ("Consistencia", "Access y CSV deben producir exactamente la misma estructura lógica."),
        ("Llaves estables", "Las llaves derivadas, como formula_key, resuelven inconsistencias reales de captura."),
        ("Caché confiable", "Streamlit necesita una huella del origen para refrescarse cuando cambian los datos."),
    ]
    for lead, text in reasons:
        add_body(doc, f"{lead}: {text}")

    add_heading(doc, "6. Resultado práctico")
    add_body(doc, "Con esta división, el proyecto queda más fácil de mantener: cambiar una ruta, una tabla o una regla de limpieza no obliga a tocar la pantalla completa.")
    add_body(doc, "También queda más fácil de probar: cada bloque se valida por separado y el dashboard solo orquesta.")

    add_heading(doc, "7. Si quieres probarlo en tu Mac")
    add_body(doc, "Primero activas el entorno virtual, luego instalas las dependencias, ejecutas streamlit run app.py y trabajas en modo DEV con los CSV de esta carpeta.")

    doc.save(OUT)


if __name__ == "__main__":
    main()
